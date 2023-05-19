import datetime
from bleach.sanitizer import Cleaner
from docx import Document
from docx.shared import Cm

WordCleaner = Cleaner(tags='', strip=True)

def export_catalogus_to_word(instrumenten, filename):
    # De export is gebaseerd op een template waarin de juiste stijlen zijn gedefinieerd. 
    # Deze template staat in de root folder van het project.
    # Standaard worden alle gebruikte HTML-tags integraal verwijderd (en dus niet vervangen door Word-opmaak).


    # BEKENDE FOUT: Als een veld een getal bevat denkt Python dat het een int is, en dat gaan Wordcleaner.clean niet goed!

    catalogus = Document('Template.docx')
    section = catalogus.sections[0]
    footer = section.footer.paragraphs[0]

    # Voorblad is alleen  nodig als de export meerdere instrrumenten bevat
    if len(instrumenten) > 1:
        catalogus.add_heading('Catalogus interventie.afm.nl', level=0)
        catalogus.add_page_break()
        footer.text = f'Catalogus van interventie.afm.nl. Document gemaakt op {datetime.date.today()}.'
    if len(instrumenten) == 1:
        #catalogus.add_heading('Interventieinstrument', level=0)
        footer.text = f'{instrumenten[0].naam} op interventie.afm.nl. Document gemaakt op {datetime.date.today()}.'

    # Per instrument wordt een hoofdstuk aangemaakt en gevuld met paragrafen.
    for instrument in instrumenten:
        catalogus.add_heading(instrument.naam, level=1)

        #catalogus.add_paragraph().add_run(WordCleaner.clean(instrument.intro)).bold=True
        catalogus.add_paragraph(WordCleaner.clean(instrument.intro), style='Intro')

        catalogus.add_heading('Wanneer te gebruiken', level=2)
        catalogus.add_paragraph(WordCleaner.clean(instrument.beschrijving))

        catalogus.add_heading('Overwegingen bij gebruik', level=2)
        catalogus.add_paragraph(WordCleaner.clean(instrument.afwegingen))

        catalogus.add_heading('Voorbeelden van projecten', level=2)
        catalogus.add_paragraph(WordCleaner.clean(instrument.voorbeelden))

        catalogus.add_heading('Links naar documentatie', level=2)
        catalogus.add_paragraph(instrument.links)

        catalogus.add_heading('Tags', level=2)
        catalogus.add_heading('Tags van toepassing op dit instrument', level=3)
        paragraph = catalogus.add_paragraph()
        for tag in instrument.tags:
            paragraph.add_run('+'+tag.naam, style='PlusTag')
            paragraph.add_run(' ')
        catalogus.add_heading('Tags die dit instrument uitsluiten', level=3)
        paragraph = catalogus.add_paragraph()
        for tag in instrument.extags:
            paragraph.add_run('-'+tag.naam, style='MinTag')
            paragraph.add_run(' ')


        catalogus.add_heading('Eigenaar', level=2)
        catalogus.add_paragraph(instrument.eigenaar)
        catalogus.add_paragraph(instrument.eigenaar_email)
        if not instrument == instrumenten[-1]:
            # Plaats geen page break na het laatste instrument.
            catalogus.add_page_break()
    catalogus.save(filename)
    return filename


def export_session_to_word(werksessie, vraagcategorieen, instrumenten, filename):
    # Titelblad
    verslag = Document('Template.docx')
    verslag.add_heading('Verslag werksessie', level=0)
    verslag.add_heading(f'{WordCleaner.clean(werksessie.naam)}', level=0)
    
    verslag.add_paragraph()
    verslag.add_paragraph(f'Datum van de werksessie: {WordCleaner.clean(werksessie.datum)}')
    verslag.add_paragraph(f'Deelnemers aan de werksessie: {WordCleaner.clean(werksessie.auteurs)}')

    verslag.add_paragraph()

    section = verslag.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = f'Verslag {WordCleaner.clean(werksessie.naam)} op {WordCleaner.clean(werksessie.datum)} op interventie.afm.nl.'
    verslag.add_page_break()

    # Informatie over de casus
    verslag.add_heading('Informatie over de casus', level=1)


    verslag.add_heading('Beoogd effect', level=2)
    verslag.add_paragraph('Met alle aanwezigen vastgesteld wat het doel is dat zij met de interventie willen bereiken. Welk gedrag is onwenselijk en moet gestopt of verminderd worden? Of: welk gedrag moet gestart of voortgezet worden? De groep heeft het volgende vastgesteld:')
    verslag.add_paragraph(WordCleaner.clean(werksessie.probleemstelling))

    verslag.add_heading('Definitieve overwegingen en instrumentselectie', level=2)
    verslag.add_paragraph('De export van de werksessie biedt aanknopingspunten voor een plan van aanpak. Hieronder staat welke instrumenten relevant zijn voor deze casus en welke overwegingen daarbij een rol spelen. Ook kan besloten worden niets te doen. De aanpak kan gebaseerd zijn op de instrumenten die door de keuzehulp zijn geprioriteerd, maar dat is niet noodzakelijk.')
    verslag.add_paragraph(WordCleaner.clean(werksessie.conclusie))
    verslag.add_page_break()

    # Ingevulde antwoorden
    verslag.add_heading('Geselecteerde antwoorden', level=1)
    verslag.add_paragraph('Hieronder volgt een overzicht van de gegeven antwoorden op de vragen.')

    for categorie in vraagcategorieen:
        verslag.add_heading(categorie.naam, level=2)
        for vraag in categorie.vragen:
            verslag.add_paragraph(vraag.naam, style='Vraag')
            for optie in vraag.opties:
                if optie in werksessie.geselecteerde_opties:
                    verslag.add_paragraph(optie.naam, style='Optie')
                    
            for motivatie in werksessie.motivaties:
                if motivatie.vraag == vraag.id:
                    verslag.add_paragraph(motivatie.motivatie, style='Motivatie')
    verslag.add_page_break()

    # Advies
    verslag.add_heading('Advies op basis van de werksessie', level=1)
    verslag.add_paragraph('Deze lijst van suggesties voor interventies is gebaseerd op de antwoorden gekozen in de sessie. In het volgende hoofdstuk staan de instrumenten verder uitgewerkt.')

    table = verslag.add_table(rows=1, cols=3, style='interventie.afm.nl')
    header = table.rows[0].cells
    header[0].text = 'Instrument'
    header[1].text = 'Score'
    header[1].width = Cm(1.0)
    header[2].text = 'Omschrijving'
    header[2].width = Cm(16.5)
    for instrument in instrumenten:
        if instrument[1] > 0:
            rij = table.add_row().cells
            rij[0].text = instrument[0].naam
            rij[1].text = str(instrument[5])
            rij[1].width = Cm(1.0)
            rij[2].text = instrument[0].intro    
            rij[2].width = Cm(16.5)

    verslag.add_page_break()

    # Instrumenten
    verslag.add_heading('Mogelijke interventies', level=1)
    verslag.add_paragraph('Hieronder staat korte omschrijving van de meest toepasselijke interventies volgens de antwoorden gegeven in de werksessie. Een volledig overzicht van alle instrumenten staat in de catalogus, die in te zien is op www.interventie.afm.nl.')
    for instrument_met_info in instrumenten:
        if instrument_met_info[6] == 2:
            instrument = instrument_met_info[0]
            # Van de instrumenten die een hoge prioriteit hebben gekregen w ordt de informatie in het verslag opgenomen.
            verslag.add_heading(instrument.naam, level=2)

            verslag.add_paragraph(WordCleaner.clean(instrument.intro), style='Intro')

            verslag.add_heading('Wanneer te gebruiken', level=3)
            verslag.add_paragraph(WordCleaner.clean(instrument.beschrijving))

            verslag.add_heading('Overwegingen bij gebruik', level=3)
            verslag.add_paragraph(WordCleaner.clean(instrument.afwegingen))

            verslag.add_heading('Voorbeelden van projecten', level=3)
            verslag.add_paragraph(WordCleaner.clean(instrument.voorbeelden))

            verslag.add_heading('Eigenaar', level=3)
            verslag.add_paragraph(instrument.eigenaar)
            verslag.add_paragraph(instrument.eigenaar_email)

            verslag.add_heading('Score', level=3)
            verslag.add_paragraph(f'Gewogen score van dit instrument op basis van de werksessie: {instrument_met_info[5]}.')
            paragraph = verslag.add_paragraph()
            for tag in instrument_met_info[2]:
                paragraph.add_run('+'+tag.naam, style='PlusTag')
                paragraph.add_run(' ')
            for tag in instrument_met_info[4]:
                paragraph.add_run('-'+tag.naam, style='MinTag')
                paragraph.add_run(' ')    

            if not instrument == instrumenten[-1]:
                # Plaats geen page break na het laatste instrument.
                verslag.add_page_break()
    verslag.save(filename)
    return filename
